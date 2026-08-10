import io
from typing import Any

import pytest
from docx import Document as DocxDocument

from sightline.assembly import (
    EMPLOYER_ORDER,
    assemble,
    generate_brief,
    render_docx,
    select_bullets,
)
from sightline.provenance import ProvenanceError


def make_bullet(
    ref: str, org: str, variants: list[str], tags: list[str],
    provenance: str = "measured", status: str = "verified", period: str = "2023-2025",
    text: str | None = None,
) -> dict[str, Any]:
    return {
        "id": int(ref.split("-")[1]),
        "ref": ref,
        "text": text or f"Did the thing described by {ref}.",
        "source_org": org,
        "source_period": period,
        "tags": tags,
        "variants": variants,
        "provenance": provenance,
        "status": status,
    }


BULLETS = [
    make_bullet("BL-001", "BEAM LEGACY GROUP", ["engineer"], ["production", "api design"]),
    make_bullet("BL-002", "BEAM LEGACY GROUP", ["engineer"], ["llm", "cost"]),
    make_bullet("BL-003", "BEAM LEGACY GROUP", ["engineer"], ["qa", "validation"]),
    make_bullet("BL-004", "BEAM LEGACY GROUP", ["engineer"], ["etl", "data pipelines"]),
    make_bullet("BL-005", "COMARKCO", ["engineer"], ["finance", "operations"], period="2020-2023"),
    make_bullet("BL-006", "COMARKCO", ["engineer"], ["compliance"], period="2020-2023"),
    make_bullet("BL-007", "COMARKCO", ["engineer"], ["scale"], period="2020-2023"),
    make_bullet("BL-008", "COMARKCO", ["leadership"], ["finance", "operations"], period="2020-2023"),
]


# ---- select_bullets ----


def test_select_bullets_filters_by_variant():
    sections = select_bullets(BULLETS, "leadership", [])
    all_refs = [b["ref"] for sec in sections for b in sec["order"]]
    assert all_refs == ["BL-008"]


def test_select_bullets_groups_in_employer_order():
    sections = select_bullets(BULLETS, "engineer", [])
    orgs = [sec["org"] for sec in sections]
    assert orgs == [o for o in EMPLOYER_ORDER if o in orgs]
    assert orgs == ["BEAM LEGACY GROUP", "COMARKCO"]


def test_select_bullets_keeps_top_scored_first():
    sections = select_bullets(BULLETS, "engineer", ["LLM cost reduction"])
    beam = next(s for s in sections if s["org"] == "BEAM LEGACY GROUP")
    assert beam["order"][0]["ref"] == "BL-002"  # matched "llm"/"cost" tags


def test_select_bullets_keeps_at_least_3_of_4():
    sections = select_bullets(BULLETS, "engineer", [])
    beam = next(s for s in sections if s["org"] == "BEAM LEGACY GROUP")
    assert len(beam["order"]) == 3
    assert len(beam["dropped"]) == 1


def test_select_bullets_excludes_retired_bullets():
    # Found live: a retired bullet was still getting selected, then failing
    # assert_shippable with a confusing "status=retired" error instead of
    # being cleanly passed over — retired is a permanent exclusion, not
    # "unreviewed" the way draft is.
    bullets = [*BULLETS, make_bullet(
        "BL-009", "BEAM LEGACY GROUP", ["engineer"], ["production", "api design"], status="retired",
    )]
    sections = select_bullets(bullets, "engineer", [])
    all_refs = [b["ref"] for sec in sections for b in sec["order"]] + \
        [b["ref"] for sec in sections for b in sec["dropped"]]
    assert "BL-009" not in all_refs


def test_select_bullets_never_modifies_bullet_text():
    sections = select_bullets(BULLETS, "engineer", ["api design"])
    for sec in sections:
        for b in sec["order"]:
            original = next(x for x in BULLETS if x["ref"] == b["ref"])
            assert b["text"] == original["text"]


# ---- render_docx ----


def test_render_docx_produces_openable_document():
    sections = select_bullets(BULLETS, "engineer", [])
    doc_bytes = render_docx("engineer", sections)
    doc = DocxDocument(io.BytesIO(doc_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "JAMES A. BEAM" in text
    assert "PROFESSIONAL EXPERIENCE" in text
    assert "BEAM LEGACY GROUP" in text


def test_render_docx_includes_selected_bullet_text_not_dropped():
    sections = select_bullets(BULLETS, "engineer", [])
    doc_bytes = render_docx("engineer", sections)
    doc = DocxDocument(io.BytesIO(doc_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    beam = next(s for s in sections if s["org"] == "BEAM LEGACY GROUP")
    for b in beam["order"]:
        assert b["text"] in text
    for b in beam["dropped"]:
        assert b["text"] not in text


# ---- assemble orchestration, with fakes ----


class FakeDB:
    def __init__(self, posting, bullets):
        self._posting = posting
        self._bullets = bullets
        self.uploaded = None
        self.inserted_variant = None
        self.events = []

    def get_posting(self, posting_id):
        return self._posting

    def get_bullets_full(self):
        return self._bullets

    def upload_document(self, bucket, path, content):
        self.uploaded = (bucket, path, content)

    def create_signed_url(self, bucket, path, expires_in=3600, download_filename=None):
        url = f"https://example.supabase.co/storage/v1/object/sign/{bucket}/{path}?token=fake"
        if download_filename:
            url += f"&download={download_filename}"
        return url

    def insert_variant(self, fields):
        self.inserted_variant = {**fields, "id": 99}
        return self.inserted_variant

    def log_event(self, entity_type, event, entity_id=None, payload=None):
        self.events.append((entity_type, event, entity_id, payload))


class FakeAnthropic:
    def __init__(self):
        self.calls = []

    def structured_call(self, **kwargs):
        self.calls.append(kwargs)
        return {"brief": "Lead with the production system."}, 0.004


def make_posting(variant="engineer", verified=True):
    provenance_status = "verified" if verified else "draft"
    bullets = [
        make_bullet("BL-101", "BEAM LEGACY GROUP", [variant], ["production"],
                    status=provenance_status),
        make_bullet("BL-102", "BEAM LEGACY GROUP", [variant], ["llm"],
                    status=provenance_status),
        make_bullet("BL-103", "BEAM LEGACY GROUP", [variant], ["qa"],
                    status=provenance_status),
    ]
    posting = {
        "id": 5,
        "title": "AI Automation Engineer",
        "companies": {"id": 1, "name": "Acme"},
        "jd_text": "Looking for someone to run production AI systems.",
        "scores": [{
            "rationale": "Strong overlap.",
            "keywords": ["production"],
            "unmet_requirements": ["domain knowledge"],
            "company_signals": ["Raised $10M"],
            "suggested_variant": variant,
        }],
    }
    return posting, bullets


def test_assemble_raises_when_posting_unscored():
    posting = {"id": 5, "title": "X", "companies": {}, "scores": []}
    db = FakeDB(posting, [])
    with pytest.raises(ValueError, match="not been scored"):
        assemble(db, FakeAnthropic(), 5)


def test_assemble_blocks_on_unverified_bullets_and_logs_it():
    posting, bullets = make_posting(verified=False)
    db = FakeDB(posting, bullets)
    with pytest.raises(ProvenanceError):
        assemble(db, FakeAnthropic(), 5)
    assert db.uploaded is None
    assert db.inserted_variant is None
    blocked = [e for e in db.events if e[1] == "assembly_blocked"]
    assert len(blocked) == 1


def test_assemble_happy_path_uploads_and_records_variant():
    posting, bullets = make_posting(verified=True)
    db = FakeDB(posting, bullets)
    anthropic = FakeAnthropic()
    result = assemble(db, anthropic, 5)

    assert db.uploaded is not None
    bucket, path, content = db.uploaded
    assert bucket == "resumes"
    assert path.startswith("5/engineer-")
    assert content[:2] == b"PK"  # docx is a zip archive

    assert db.inserted_variant["kind"] == "engineer"
    assert set(db.inserted_variant["bullet_refs"]) == {"BL-101", "BL-102", "BL-103"}
    assert db.inserted_variant["brief"] == "Lead with the production system."
    assert result["signed_url"].startswith("https://")
    assert any(e[1] == "assembled" for e in db.events)

    assert result["sections"][0]["org"] == "BEAM LEGACY GROUP"


def test_assemble_returns_jd_alignment_context():
    posting, bullets = make_posting(verified=True)
    db = FakeDB(posting, bullets)
    result = assemble(db, FakeAnthropic(), 5)

    assert result["jd_text"] == "Looking for someone to run production AI systems."
    assert result["jd_keywords"] == ["production"]
    assert result["rationale"] == "Strong overlap."

    order = {b["ref"]: b for b in result["sections"][0]["order"]}
    assert order["BL-101"]["matched_keywords"] == ["production"]  # tag "production" hits the JD keyword
    assert order["BL-102"]["matched_keywords"] == []  # tag "llm" doesn't
    assert order["BL-103"]["matched_keywords"] == []  # tag "qa" doesn't
    assert {b["ref"] for b in result["sections"][0]["order"]} == {"BL-101", "BL-102", "BL-103"}


def test_assemble_respects_explicit_variant_override():
    posting, engineer_bullets = make_posting(variant="engineer", verified=True)
    _, leadership_bullets = make_posting(variant="leadership", verified=True)
    db = FakeDB(posting, engineer_bullets + leadership_bullets)
    result = assemble(db, FakeAnthropic(), 5, variant="leadership")
    assert result["kind"] == "leadership"


def test_generate_brief_sends_rationale_and_keywords_to_anthropic():
    posting, _ = make_posting()
    score = posting["scores"][0]
    anthropic = FakeAnthropic()
    brief, cost = generate_brief(anthropic, posting, score)
    assert brief == "Lead with the production system."
    assert cost == 0.004
    sent = anthropic.calls[0]["user_content"]
    assert "Strong overlap." in sent
    assert "production" in sent


def test_generate_brief_survives_missing_brief_field():
    # BRIEF_SCHEMA marks "brief" required, but tool-use generation
    # completeness isn't guaranteed — found live: assembly's first real
    # invocation all session (blocked by the provenance gate until then)
    # hit exactly this and crashed with a bare KeyError.
    class EmptyAnthropic:
        def structured_call(self, **kwargs):
            return {}, 0.002

    posting, _ = make_posting()
    score = posting["scores"][0]
    brief, cost = generate_brief(EmptyAnthropic(), posting, score)
    assert brief == ""
    assert cost == 0.002
