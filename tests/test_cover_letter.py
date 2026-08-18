from unittest.mock import MagicMock

from docx import Document as DocxDocument

from sightline.cover_letter import (
    AI_PROJECTS,
    build_user_content,
    generate_cover_letter,
    greeting_for,
    render_cover_letter_docx,
)
from sightline.voice import COVER_LETTER_EXAMPLES, VOICE_RULES

BULLETS = [
    {"ref": "BL-001", "text": "Built the automation platform.", "status": "verified"},
    {"ref": "BL-002", "text": "Scaled revenue from $250K to $12M.", "status": "verified"},
    {"ref": "BL-003", "text": "Unverified claim.", "status": "draft"},
]

POSTING = {
    "title": "Program Operations Manager",
    "companies": {"name": "Convergent Research"},
    "jd_text": "Own applicant pipeline operations end to end.",
}

SCORE = {
    "rationale": "Strong overlap.",
    "keywords": ["applicant pipeline"],
    "unmet_requirements": ["startup ops background"],
    "company_signals": ["nonprofit research studio"],
}


def test_generate_cover_letter_raises_on_empty_response():
    # Found live: extended thinking consumed the whole token budget and the
    # model returned empty text with stop_reason=max_tokens — that was
    # silently saved as the cover letter and uploaded to storage. Must raise
    # instead of returning garbage.
    fake_client = MagicMock()
    fake_client.chat_call.return_value = ("", 0.02)
    try:
        generate_cover_letter(fake_client, POSTING, SCORE, BULLETS, ["BL-001"])
        assert False, "expected ValueError"
    except ValueError as e:
        assert "no usable text" in str(e)


def test_build_user_content_includes_jd_and_score_context():
    content = build_user_content(POSTING, SCORE)
    assert "Program Operations Manager" in content
    assert "Convergent Research" in content
    assert "Own applicant pipeline operations end to end." in content
    assert "Strong overlap." in content


def test_generate_cover_letter_only_grounds_in_verified_bullets():
    fake_client = MagicMock()
    fake_client.chat_call.return_value = (
        "Paragraph one, with enough real content to clear the length floor.\n\nParagraph two.", 0.015,
    )
    text, cost = generate_cover_letter(fake_client, POSTING, SCORE, BULLETS, ["BL-001"])

    # No fixed closing appended as of the writing-guide rewrite — the
    # model's own output is returned as-is (see the "as-is" test below for
    # the direct regression coverage of that).
    assert text == (
        "Paragraph one, with enough real content to clear the length floor.\n\nParagraph two."
    )
    assert cost == 0.015
    sent_system = fake_client.chat_call.call_args.kwargs["system"]
    assert "Built the automation platform." in sent_system
    assert "Scaled revenue from $250K to $12M." in sent_system
    assert "Unverified claim." not in sent_system


def test_generate_cover_letter_includes_voice_rules():
    fake_client = MagicMock()
    fake_client.chat_call.return_value = ("xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx", 0.01)
    generate_cover_letter(fake_client, POSTING, SCORE, BULLETS, ["BL-001"])
    sent_system = fake_client.chat_call.call_args.kwargs["system"]
    assert VOICE_RULES in sent_system
    assert "spearheaded" in sent_system  # banned-word list present
    assert "em dash" in sent_system


def test_generate_cover_letter_includes_voice_reference_from_answers():
    answers = [
        {"question_type": "tell_me_about_a_failure", "status": "ready",
         "text": "I'm using that word generously — " + ("x" * 220)},
        {"question_type": "salary_expectations", "status": "ready", "text": "short one"},  # too short to use
        {"question_type": "biggest_system_built", "status": "draft",
         "text": "y" * 300},  # draft — never used, even though it's long
    ]
    fake_client = MagicMock()
    fake_client.chat_call.return_value = ("xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx", 0.01)
    generate_cover_letter(fake_client, POSTING, SCORE, BULLETS, ["BL-001"], answers)
    sent_system = fake_client.chat_call.call_args.kwargs["system"]
    assert "I'm using that word generously" in sent_system
    assert "short one" not in sent_system
    assert "y" * 300 not in sent_system


def test_generate_cover_letter_works_without_answers():
    fake_client = MagicMock()
    fake_client.chat_call.return_value = ("xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx", 0.01)
    text, cost = generate_cover_letter(fake_client, POSTING, SCORE, BULLETS, ["BL-001"])
    assert text == "xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx"
    sent_system = fake_client.chat_call.call_args.kwargs["system"]
    assert "(none available)" in sent_system


def test_generate_cover_letter_returns_model_text_as_is_no_fixed_closing():
    # The close used to be a fixed string appended after generation, to
    # avoid the model inventing a follow-up commitment ("I'll follow up
    # next week") the candidate has no intention of keeping. The 2026-08-17
    # writing guide's confirmed real letters each end differently — a fixed
    # line traded away a genuine, letter-specific close, so the prompt now
    # steers away from invented commitments directly instead (see
    # SYSTEM_PROMPT) and the model's own output is returned untouched.
    fake_client = MagicMock()
    fake_client.chat_call.return_value = ("Body text with enough length to clear the floor check.", 0.01)
    text, _ = generate_cover_letter(fake_client, POSTING, SCORE, BULLETS, ["BL-001"])
    assert text == "Body text with enough length to clear the floor check."


def test_generate_cover_letter_includes_ai_projects_and_confirmed_examples():
    fake_client = MagicMock()
    fake_client.chat_call.return_value = ("xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx", 0.01)
    generate_cover_letter(fake_client, POSTING, SCORE, BULLETS, ["BL-001"])
    sent_system = fake_client.chat_call.call_args.kwargs["system"]
    assert AI_PROJECTS in sent_system
    assert COVER_LETTER_EXAMPLES in sent_system
    assert "**Sightline**" in sent_system


def test_generate_cover_letter_uses_new_opener_format():
    fake_client = MagicMock()
    fake_client.chat_call.return_value = ("xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx", 0.01)
    generate_cover_letter(fake_client, POSTING, SCORE, BULLETS, ["BL-001"])
    sent_system = fake_client.chat_call.call_args.kwargs["system"]
    assert "I'm excited to apply for the [exact role name] role at [Company]." in sent_system


def test_generate_cover_letter_separates_selected_from_other_bullets():
    fake_client = MagicMock()
    fake_client.chat_call.return_value = ("xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx", 0.01)
    generate_cover_letter(fake_client, POSTING, SCORE, BULLETS, ["BL-001"])
    sent_system = fake_client.chat_call.call_args.kwargs["system"]
    # BL-001 (selected) should appear before the "rest of the library" section
    selected_idx = sent_system.index("Built the automation platform.")
    other_idx = sent_system.index("Scaled revenue from $250K to $12M.")
    assert selected_idx < other_idx


def test_render_cover_letter_docx_produces_valid_docx_with_body_text():
    docx_bytes = render_cover_letter_docx(
        "First paragraph of the letter.\n\nSecond paragraph.", "Convergent Research", "Program Operations Manager",
    )
    assert docx_bytes[:2] == b"PK"  # docx is a zip archive
    import io
    doc = DocxDocument(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "First paragraph of the letter." in full_text
    assert "Second paragraph." in full_text
    assert "Convergent Research" in full_text
    assert "Program Operations Manager" in full_text
    assert "Dear Hiring Team," in full_text  # default greeting


def test_render_cover_letter_docx_uses_best_sign_off():
    # All four confirmed real letters (CodePath, Mercury, Argano, Whip
    # Around) sign off "Best," — not "Sincerely,".
    docx_bytes = render_cover_letter_docx("Body.", "Acme", "Role")
    import io
    doc = DocxDocument(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Best," in full_text
    assert "Sincerely," not in full_text


def test_render_cover_letter_docx_bolds_ai_project_name_in_bullets():
    docx_bytes = render_cover_letter_docx(
        "Intro paragraph.\n\n- **Sightline** Built and launched a workflow system.\n- "
        "**Barometer** Built and tested a prediction market system.\n\nClosing paragraph.",
        "Acme", "Role",
    )
    import io
    doc = DocxDocument(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # The markdown markers themselves never survive into the document.
    assert "**" not in full_text
    assert "Sightline" in full_text
    assert "Built and launched a workflow system." in full_text
    bullet_para = next(p for p in doc.paragraphs if p.text.startswith("Sightline"))
    assert bullet_para.runs[0].text == "Sightline"
    assert bullet_para.runs[0].bold is True
    assert bullet_para.runs[1].bold is not True


def test_render_cover_letter_docx_uses_custom_greeting():
    docx_bytes = render_cover_letter_docx("Body.", "Acme", "Role", greeting="Dear Kait Picco,")
    import io
    doc = DocxDocument(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Dear Kait Picco," in full_text
    assert "Dear Hiring Team," not in full_text


def test_greeting_for_uses_named_contact_when_present():
    score = {"named_contacts": [{"name": "Kait Picco", "title": "Head of People", "context": "..."}]}
    assert greeting_for(score) == "Dear Kait Picco,"


def test_greeting_for_falls_back_when_no_named_contact():
    assert greeting_for({}) == "Dear Hiring Team,"
    assert greeting_for({"named_contacts": []}) == "Dear Hiring Team,"
    assert greeting_for({"named_contacts": [{"name": ""}]}) == "Dear Hiring Team,"
