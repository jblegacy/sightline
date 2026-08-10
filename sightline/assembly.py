"""Stage 4 — Assemble. See docs/SIGHTLINE_BUILD_SPEC_V2.md §7 and CLAUDE.md
rule 1: selection only, never generate new bullet text.

The provenance validator (sightline/provenance.py) gates every bullet that
goes into a rendered document — this module never catches ProvenanceError to
continue past it, only to log the block before re-raising.
"""
from __future__ import annotations

import io
import math
import re
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from sightline.anthropic_client import AnthropicClient
from sightline.db import SightlineDB
from sightline.provenance import Bullet, ProvenanceError, assert_shippable

STORAGE_BUCKET = "resumes"
BRIEF_MODEL = "claude-sonnet-5"

# ---- pre-approved static content ----
#
# This is the user's own already-approved resume text (see
# scripts/build_resumes.js, the source of truth this was ported from) — not
# a claim pulled from the `bullets` library, so it isn't run through the
# provenance validator. Only the PROFESSIONAL EXPERIENCE bullets below are
# selected per posting; everything here is fixed across every build.

NAME = "JAMES A. BEAM, MBA"
CONTACT = (
    "Albany, OR  |  714-322-9914  |  james@beamlegacy.com  |  "
    "linkedin.com/in/jamesabeam  |  beamlegacy.com"
)

# What actually lands on disk when downloaded — the storage object key below
# keeps the variant/timestamp for internal uniqueness, but nobody applying
# for a job wants a file named "leadership-20260805T080412Z.docx". The
# variant type is surfaced separately in the dashboard as a tag, not baked
# into the filename.
_FILENAME_UNSAFE = re.compile(r'[/\\:*?"<>|]')


def resume_download_name(company: str) -> str:
    company = _FILENAME_UNSAFE.sub("", company).strip() or "Unknown"
    return f"James Beam - Resume - {company}.docx"

EDUCATION = [
    "Master of Business Administration — Concordia University Irvine, 2017",
    "Bachelor of Science, Business Administration (Entrepreneurship) — Chapman University, 2011",
]

VARIANT_CONTENT: dict[str, dict[str, Any]] = {
    "leadership": {
        "headline": "AI Enablement & Operations  |  Workflow Automation  |  Business Operations",
        "summary": (
            "Operations leader who designs, builds, and ships AI-automated business systems. "
            "Scaled a consumer packaged goods brand from $250K to $12M in twelve months by "
            "building the operational, financial, and compliance infrastructure behind it, then "
            "built and shipped AI-native software platforms from scratch — including an "
            "autonomous decision-automation system running live in production. Combines 12+ "
            "years of cross-functional operations leadership with hands-on technical "
            "implementation: agentic workflow design, workflow redesign, process automation, "
            "adoption and change management, and measurable time-and-cost reduction. MBA. "
            "Remote-first operator."
        ),
        "competencies": (
            "AI Enablement & Adoption  •  AI Workflow Automation  •  Process Optimization"
            "  •  Business Operations  •  Cross-Functional Collaboration  •  Change "
            "Management  •  Program & Project Management  •  Product Ownership  •  "
            "Data & Business Intelligence  •  Financial Planning & Analysis  •  "
            "Operational Scaling  •  SOP & Playbook Design  •  Stakeholder & Vendor "
            "Management  •  Team Leadership"
        ),
        "skills": [
            (
                "AI & Automation",
                "Anthropic Claude API (Batch API, prompt caching, Model Context Protocol), "
                "Claude Code, multi-agent orchestration, agentic workflow design, prompt "
                "engineering, LLM evaluation, n8n, Make, Zapier, Microsoft Power Automate",
            ),
            (
                "Data & Development",
                "Python, SQL, REST API integration, Flask, React, JavaScript, Git / GitHub, "
                "Railway, data pipelines, vector embeddings",
            ),
            (
                "Business Systems",
                "QuickBooks, NetSuite, Shopify, Klaviyo, Asana, Slack, Google Workspace, "
                "Microsoft 365, Excel / Google Sheets modeling",
            ),
            (
                "Methodologies",
                "Agile, Scrum, PMO governance, Sales & Operations Planning (S&OP), process "
                "mapping, statistical validation (walk-forward, Monte Carlo)",
            ),
        ],
    },
    "engineer": {
        "headline": "Workflow Automation & Business Systems  |  AI Enablement  |  Operations Technology",
        "summary": (
            "Operations leader who automates the work. Twelve years running business "
            "operations — finance, logistics, compliance, reporting — now paired with "
            "hands-on AI tooling: specifies and ships production automation by directing AI "
            "coding agents, including an autonomous decision-automation platform running live "
            "today. Builds against how a business actually runs rather than against an "
            "idealized process diagram. Seeking hands-on individual-contributor work rather "
            "than people management. Remote-first."
        ),
        "competencies": (
            "Workflow Automation  •  Business Process Automation  •  Systems "
            "Integration  •  AI Enablement & Adoption  •  Internal Tooling  •  "
            "Requirements Analysis  •  Business Process Mapping  •  Data & "
            "Reporting  •  Change Management  •  Vendor & Stakeholder Management  •  "
            "SOP Design  •  Operational Scaling"
        ),
        "skills": [
            (
                "Automation Platforms",
                "n8n, Zapier, Make, Microsoft Power Automate, workflow design and "
                "orchestration, system-to-system integration",
            ),
            (
                "AI Tooling",
                "Anthropic Claude API, Claude Code, agentic workflow design, prompt "
                "engineering, AI-assisted process redesign, LLM output validation",
            ),
            (
                "Data & Systems",
                "SQL, Python scripting, REST API integration, webhooks, data pipelines, "
                "reporting and dashboard design, spreadsheet modeling",
            ),
            (
                "Business Systems",
                "QuickBooks, NetSuite, Shopify, Klaviyo, Asana, Slack, Google Workspace, "
                "Microsoft 365",
            ),
            (
                "Practices",
                "Business process mapping, requirements analysis, SOP and playbook design, "
                "Agile / Scrum delivery, technical documentation, Sales & Operations "
                "Planning (S&OP)",
            ),
        ],
    },
}

# Employer order matches the master resume (scripts/build_resumes.js) and the
# seed order in supabase/migrations/0002_bullets_answers.sql — most recent
# first. Location/title/scope are fixed per variant; dates come from the
# bullets themselves (source_period) rather than being duplicated here.
EMPLOYER_ORDER = [
    "BEAM LEGACY GROUP",
    "COMARKCO",
    "WORKSITE LABS",
    "STAIRCASE DIGITAL / WHEELHOUSE STUDIO",
    "UNIVERSITY OF CALIFORNIA, IRVINE",
]

EMPLOYER_META: dict[str, dict[str, Any]] = {
    "BEAM LEGACY GROUP": {
        "location": "Albany, OR (Remote)",
        "leadership": {
            "title": "Founder & Principal Operator",
            "scope": (
                "Bootstrapped holding company building AI-native ventures. Specify, build, "
                "and ship production software using AI coding agents — operating as "
                "product owner, operator, and implementer across the portfolio."
            ),
        },
        "engineer": {
            "title": "Founder — Operator & Principal Implementer",
            "scope": (
                "Bootstrapped venture studio. Hands-on across specification, implementation, "
                "deployment, and day-to-day operation of production AI systems — built by "
                "directing AI coding agents rather than a development team."
            ),
        },
    },
    "COMARKCO": {
        "location": "Los Angeles, CA",
        "leadership": {
            "title": "Chief of Staff  |  Fractional COO / CFO",
            "scope": (
                "Served as fractional COO/CFO across a portfolio of consumer packaged goods "
                "(CPG) brands, building a repeatable operating model covering customer "
                "experience, logistics, HR, IT, legal and compliance, data analytics, "
                "bookkeeping, and tax."
            ),
        },
        "engineer": {
            "title": "Chief of Staff  |  Fractional Operations & Finance Lead",
            "scope": (
                "Built and automated the operating systems for a portfolio of consumer "
                "packaged goods brands — finance, logistics, compliance, reporting, and "
                "data analytics."
            ),
        },
    },
    "WORKSITE LABS": {
        "location": "Los Angeles, CA",
        "leadership": {
            "title": "SVP of Experience & Product",
            "scope": (
                "Led customer experience and product strategy through hyper-growth from $0 "
                "to $25M in revenue in year one."
            ),
        },
        "engineer": {
            "title": "SVP of Experience & Product",
            "scope": (
                "Product owner and systems builder through growth from $0 to $25M in revenue "
                "in year one."
            ),
        },
    },
    "STAIRCASE DIGITAL / WHEELHOUSE STUDIO": {
        "location": "Los Angeles, CA",
        "leadership": {
            "title": "Chief of Staff  |  Director of Operations",
            "scope": (
                "Strategic partner to executives across operations, fundraising, and digital "
                "product; led development of MyIPO, a Reg A+ digital securities offering "
                "platform, and Next Ones, an athlete discovery platform."
            ),
        },
        "engineer": {
            "title": "Director of Operations  |  Product & Delivery Lead",
            "scope": (
                "Led delivery of MyIPO, a Reg A+ digital securities offering platform, and "
                "Next Ones, an athlete discovery platform."
            ),
        },
    },
    "UNIVERSITY OF CALIFORNIA, IRVINE": {
        "location": "Irvine, CA",
        "leadership": {"title": "Operations Analyst, Civil & Environmental Engineering", "scope": None},
        "engineer": {"title": "Operations Analyst, Civil & Environmental Engineering", "scope": None},
    },
}


def _format_dates(period: str) -> str:
    if "-" in period:
        start, end = period.split("-", 1)
        return f"{start} – {end}"
    return period


def select_bullets(
    bullets: list[dict[str, Any]], variant: str, keywords: list[str]
) -> list[dict[str, Any]]:
    """Group bullets by employer in resume order, score each by tag overlap
    against the JD's extracted keywords, and keep the top ~70% (never fewer
    than 3) per employer section, highest relevance first. Ported from
    prototype/sightline-dashboard.html's selectBullets() — same algorithm,
    real data. Selection only; bullet text is never modified."""
    kw = [(k.lower(), k) for k in keywords]  # (comparison form, original casing)

    def matched(tags: list[str]) -> list[str]:
        # which JD keywords this bullet's tags actually hit — the real,
        # mechanical "why" behind keep/drop, not an invented explanation.
        # Returned in the JD's own casing, matching the keyword chips shown
        # alongside them.
        out, seen = [], set()
        for t in tags:
            tl = t.lower()
            for k_lower, k_orig in kw:
                if (k_lower in tl or tl in k_lower) and k_lower not in seen:
                    out.append(k_orig)
                    seen.add(k_lower)
        return out

    by_org: dict[str, list[dict[str, Any]]] = {}
    for b in bullets:
        # retired is a permanent exclusion, not "unreviewed" — a retired
        # bullet should never be a candidate, not just fail assert_shippable
        # after already being selected (found live: it surfaced as a
        # confusing assembly error instead of being cleanly passed over).
        if b.get("status") == "retired":
            continue
        if variant in (b.get("variants") or []) and b.get("source_org") in EMPLOYER_META:
            by_org.setdefault(b["source_org"], []).append(b)

    sections = []
    for org in EMPLOYER_ORDER:
        group = by_org.get(org)
        if not group:
            continue
        scored = [(matched(b.get("tags") or []), i, b) for i, b in enumerate(group)]
        keep_n = max(3, math.ceil(len(group) * 0.7))
        ranked = sorted(scored, key=lambda x: (-len(x[0]), x[1]))
        kept, dropped = ranked[:keep_n], ranked[keep_n:]

        def with_match(item: tuple[list[str], int, dict[str, Any]]) -> dict[str, Any]:
            match_kw, _, b = item
            return {**b, "matched_keywords": match_kw}

        sections.append({
            "org": org,
            "order": [with_match(item) for item in kept],
            "dropped": [with_match(item) for item in sorted(dropped, key=lambda x: x[1])],
        })
    return sections


def render_docx(variant: str, sections: list[dict[str, Any]]) -> bytes:
    content = VARIANT_CONTENT[variant]
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    # 10pt / 0.5in margins packed more onto a page but sit outside what 2026
    # ATS-formatting guidance calls safe (11-12pt body, 0.75-1in margins) —
    # tight margins and small type are what get flagged, not page count (1-2
    # pages is explicitly fine for someone with real experience).
    style.font.size = Pt(11)
    for s in doc.sections:
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.75)

    def para(text: str, *, size: float, bold: bool = False, italic: bool = False,
              before: float = 0, after: float = 5) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic

    def section_header(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(11)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(11)
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "444444")
        pbdr.append(bottom)
        p._p.get_or_add_pPr().append(pbdr)

    def skill_line(label: str, text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3.5)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)

    def employer_line(org: str, location: str, dates: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
        r1 = p.add_run(org)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(f"  |  {location}")
        r2.font.size = Pt(11)
        r3 = p.add_run(f"\t{dates}")
        r3.font.size = Pt(11)

    def bullet_line(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.size = Pt(11)

    para(NAME, size=16, bold=True, after=2)
    para(CONTACT, size=10, after=8)
    para(content["headline"], size=11, bold=True, after=8)

    section_header("PROFESSIONAL SUMMARY")
    para(content["summary"], size=11)

    section_header("CORE COMPETENCIES")
    para(content["competencies"], size=11)

    section_header("TECHNICAL SKILLS")
    for label, text in content["skills"]:
        skill_line(label, text)

    section_header("PROFESSIONAL EXPERIENCE")
    for sec in sections:
        meta = EMPLOYER_META[sec["org"]]
        variant_meta = meta[variant]
        dates = _format_dates(sec["order"][0]["source_period"]) if sec["order"] else ""
        employer_line(sec["org"], meta["location"], dates)
        para(variant_meta["title"], size=11, bold=True, italic=True, after=3)
        if variant_meta.get("scope"):
            para(variant_meta["scope"], size=11, after=4)
        for b in sec["order"]:
            bullet_line(b["text"])

    section_header("EDUCATION")
    for line in EDUCATION:
        para(line, size=11, after=2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


BRIEF_SYSTEM_PROMPT = """You write a one-page prep brief for a job candidate about to apply to \
a posting. This is private prep material only the candidate reads before a screen — you are \
not writing resume or outreach content.

Synthesize, in your own words, from the score rationale, keywords, gaps, and company signals \
given: why this posting is a good match, which keywords to naturally mirror in conversation, \
what gaps to be ready to address, and up to three specific talking points for this JD. Write it \
as 3-5 tight connected sentences, not a bulleted list."""

BRIEF_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {"brief": {"type": "string"}},
    "required": ["brief"],
}


def generate_brief(
    client: AnthropicClient, posting: dict[str, Any], score: dict[str, Any]
) -> tuple[str, float]:
    company = (posting.get("companies") or {}).get("name", "this company")
    user_content = f"""Posting: {posting.get('title')} at {company}

Score rationale: {score.get('rationale') or ''}
Keywords to mirror: {', '.join(score.get('keywords') or [])}
Gaps to prepare for: {', '.join(score.get('unmet_requirements') or [])}
Company signals: {', '.join(score.get('company_signals') or [])}"""
    result, cost_usd = client.structured_call(
        model=BRIEF_MODEL,
        system=BRIEF_SYSTEM_PROMPT,
        user_content=user_content,
        tool_name="submit_brief",
        tool_description="Submit the one-page prep brief.",
        input_schema=BRIEF_SCHEMA,
        max_tokens=800,
    )
    # BRIEF_SCHEMA marks "brief" required, but tool-use generation
    # completeness isn't guaranteed (same class of bug found live in
    # scoring.py) — brief is private prep material, not a shipped claim, so
    # falling back to empty is safe; it must never block the actual .docx.
    return result.get("brief") or "", cost_usd


def assemble(
    db: SightlineDB,
    anthropic: AnthropicClient,
    posting_id: int,
    variant: str | None = None,
) -> dict[str, Any]:
    """Orchestrates Stage 4 for one posting: select bullets, gate on
    provenance, render the .docx, upload it, generate the brief, and record
    the variant. Raises ProvenanceError (not caught here) if any selected
    bullet isn't shippable — assembly stops, it does not degrade."""
    posting = db.get_posting(posting_id)
    scores = posting.get("scores") or []
    if not scores:
        raise ValueError(f"posting {posting_id} has not been scored yet")
    score = scores[0]

    variant = variant or score.get("suggested_variant") or "engineer"
    if variant not in VARIANT_CONTENT:
        raise ValueError(f"unknown variant {variant!r}")

    bullets = db.get_bullets_full()
    sections = select_bullets(bullets, variant, score.get("keywords") or [])
    selected = [b for sec in sections for b in sec["order"]]
    if not selected:
        raise ValueError("no bullets matched this variant — nothing to assemble")

    try:
        assert_shippable(
            Bullet(ref=b["ref"], text=b["text"], provenance=b["provenance"], status=b["status"])
            for b in selected
        )
    except ProvenanceError as e:
        db.log_event(
            entity_type="posting", entity_id=posting_id, event="assembly_blocked",
            payload={"variant": variant, "reason": str(e)},
        )
        raise

    docx_bytes = render_docx(variant, sections)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    path = f"{posting_id}/{variant}-{timestamp}.docx"
    db.upload_document(STORAGE_BUCKET, path, docx_bytes)

    brief, brief_cost = generate_brief(anthropic, posting, score)

    variant_row = db.insert_variant({
        "posting_id": posting_id,
        "kind": variant,
        "bullet_refs": [b["ref"] for b in selected],
        "summary_key": f"{variant}-master",
        "storage_path": path,
        "brief": brief,
    })
    db.log_event(
        entity_type="variant", entity_id=variant_row["id"], event="assembled",
        payload={
            "posting_id": posting_id,
            "variant": variant,
            "bullet_count": len(selected),
            "brief_cost_usd": round(brief_cost, 5),
        },
    )
    company_name = (posting.get("companies") or {}).get("name", "Unknown")
    signed_url = db.create_signed_url(
        STORAGE_BUCKET, path, download_filename=resume_download_name(company_name)
    )
    return {
        **variant_row, "signed_url": signed_url, "sections": _serialize_sections(sections),
        "jd_text": posting.get("jd_text"), "jd_keywords": score.get("keywords") or [],
        "rationale": score.get("rationale") or "",
    }


def _serialize_sections(sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "org": sec["org"],
            "order": [
                {"ref": b["ref"], "text": b["text"], "matched_keywords": b.get("matched_keywords") or []}
                for b in sec["order"]
            ],
            "dropped": [
                {"ref": b["ref"], "text": b["text"], "matched_keywords": b.get("matched_keywords") or []}
                for b in sec["dropped"]
            ],
        }
        for sec in sections
    ]


def variant_detail(db: SightlineDB, posting_id: int) -> dict[str, Any]:
    """Read-only: recomputes the same selection view (pure, free) for an
    already-assembled posting and mints a fresh signed URL. Used to restore
    the diff view and download link after a page reload without re-running
    Sonnet or re-uploading a new document — a fresh assemble() would also
    cost money and create a second Storage object for no reason."""
    posting = db.get_posting(posting_id)
    variants = posting.get("variants") or []
    if not variants:
        raise LookupError(f"posting {posting_id} has no assembled variant")
    variant_row = variants[0]
    scores = posting.get("scores") or []
    score = scores[0] if scores else {}

    bullets = db.get_bullets_full()
    sections = select_bullets(bullets, variant_row["kind"], score.get("keywords") or [])
    company_name = (posting.get("companies") or {}).get("name", "Unknown")
    signed_url = db.create_signed_url(
        STORAGE_BUCKET, variant_row["storage_path"], download_filename=resume_download_name(company_name)
    )
    cover_letter_signed_url = None
    if variant_row.get("cover_letter_storage_path"):
        cover_letter_signed_url = db.create_signed_url(STORAGE_BUCKET, variant_row["cover_letter_storage_path"])
    return {
        **variant_row, "signed_url": signed_url, "sections": _serialize_sections(sections),
        "jd_text": posting.get("jd_text"), "jd_keywords": score.get("keywords") or [],
        "rationale": score.get("rationale") or "",
        "cover_letter_signed_url": cover_letter_signed_url,
    }
