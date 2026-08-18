"""Cover letter generation — prose, not bullet selection, so this is a
different category from resume assembly (CLAUDE.md rule 1 is about resume
bullet text specifically). Same discipline as outreach drafts and the
answer workbench: grounded only in verified bullets and the posting's real
details, never inventing an achievement.

Structure and voice follow James Beam's Job Application Writing Guide
(2026-08-17) and the four real cover letters confirmed alongside it
(CodePath, Mercury, Argano, Whip Around) — see sightline/voice.py
COVER_LETTER_EXAMPLES for two given in full. That guide describes ONE
canonical approach, not several formats to pick from; STYLE_STRUCTURES below
keeps three length variants of it (a full call-out-everything version, a
shorter one, and a compressed note) rather than collapsing to a single
shape, so the sandbox still gives a real choice — but all three now share
the same opener, the same AI-projects bullet treatment, and the same
model-written (not fixed) closing.
"""
from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.shared import Inches, Pt

from sightline.anthropic_client import AnthropicClient
from sightline.assembly import CONTACT, NAME
from sightline.voice import COVER_LETTER_EXAMPLES, VOICE_RULES, voice_reference

MODEL = "claude-sonnet-5"

# Pre-approved, fixed text for the three AI projects — same principle as
# resume bullets: the numbers (48 hours, 50 roles/day, 540 profiles, 60
# days) must stay exact, not get paraphrased into something subtly wrong.
# The model selects which of these to use and may lightly smooth the
# sentence for flow, but the facts inside each description are fixed. Format
# matches the confirmed real letters exactly: "- **Name** description," no
# colon (VOICE RULE 1 bans colons in prose) — **Name** is the one narrow,
# explicit exception to VOICE RULE 11's no-markdown rule, since it's the
# real letters' own convention and render_cover_letter_docx() parses it into
# an actual bold run, not literal asterisks in the delivered document.
AI_PROJECTS = """- **Sightline** Built and launched an AI powered workflow system in 48 hours \
that handles job discovery, qualification, document generation, application tracking, outreach, \
and follow ups while keeping human review in the loop. It now surfaces roughly 50 roles a day and \
reduced a two hour daily search process to a morning review.
- **Barometer** Built and tested an automated prediction market system using live market and \
weather APIs and a 540 profile testing matrix across hundreds of model configurations. The \
testing showed that market speed and weather data latency removed the trading edge, so paused \
the system rather than continuing to automate something that did not make economic sense.
- **Ottonimus** Built a functional AI powered SEO and GEO platform prototype over roughly 60 \
days, bringing website performance, technical SEO, backlinks, local SEO, Google presence, and \
emerging GEO workflows into one system."""

STYLE_STRUCTURES = {
    "traditional": """Structure, in this order — the fullest of the three formats, HARD CAP 450 \
words total (the four confirmed real examples run 380-420 words each; this format allows a bit \
more room but 450 is still a ceiling, not a suggestion):
1. Opening — the shared opener below. Nothing else in that sentence.
2. ONE short paragraph (2-3 sentences) on why this specific role caught his attention — the \
actual problem the company is solving, connected to the kind of work he likes doing. Never \
restate the JD back ("You are looking for someone to build AI-powered operational systems") — say \
something more natural, like naming what the work really is underneath the posting's own words.
3. The AI-projects bullet list, using 2-3 of the projects given below (not necessarily all \
three — pick whichever are most relevant to this role; never force all three in if only two fit). \
Format each as "- Name description" using the given text close to verbatim — don't pad or \
elaborate on it.
4. ONE paragraph (2-3 sentences) reflecting on what building these actually taught him — not \
that he knows AI tools, but the judgment part: deciding what should be automated, where people \
still need to make the call, how to make something reliable enough to be trusted. This is the \
"AI judgment, not just AI building" beat.
5. ONE paragraph (2-3 sentences) on his broader operating background (the Comarkco $250K to $12M \
story, or the inventory-visibility example if the role is about operational transformation \
specifically) — proof this approach predates AI, not proof he can recite a resume.
6. ONE paragraph (1-2 sentences) connecting specifically to this company — why this company and \
role make sense for him, referencing something real and specific about them, not "great fit" \
language.
7. ONE closing sentence, written fresh for this letter, in the spirit of "I'd love the \
opportunity to bring [something specific] to [Company] and [do something specific]" — never a \
generic "I believe I would be a great fit" close, and never invented follow-up commitments \
("I'll follow up next week").""",
    "compressed": """Structure, in this order — the shortest of the three formats, target \
150-220 words total, no padding:
1. Opening — the shared opener below.
2. ONE tight paragraph: why the role caught his attention, connected directly into his broader \
operating background in the same paragraph rather than as a separate section.
3. ONE or two lines from the AI-projects bullet list below — pick the single most relevant \
project, or two only if both are clearly relevant. Do not use all three; that's what the fuller \
formats are for.
4. ONE closing sentence, written fresh, short and direct.
This format skips the reflection paragraph and the company-specific paragraph — length discipline \
is the point, not full coverage.""",
    "warm": """Structure, in this order — HARD CAP 400 words total (the four confirmed real \
examples run 380-420 words each; treat 400 as a ceiling, not a suggestion), matching the shape of \
the confirmed real examples most closely:
1. Opening — the shared opener below. Nothing else in that sentence.
2. ONE short paragraph (2-3 sentences) on why this specific role caught his attention — the real \
problem the company is solving, connected to the kind of work he likes. Never restate the JD back \
at the reader.
3. The AI-projects bullet list, using 2-3 of the projects given below, whichever are most \
relevant to this role. Format each as "- Name description," using the given text close to \
verbatim — don't pad or elaborate on it.
4. ONE paragraph (2-3 sentences) reflecting on what building these taught him — the judgment \
part, not the tool part.
5. ONE paragraph (2-3 sentences) on his broader operating background (Comarkco, or the inventory \
example if the role is about operational transformation specifically).
6. ONE paragraph (1-2 sentences) connecting specifically to this company.
7. ONE closing sentence, written fresh, in the spirit of "I'd love the opportunity to bring \
[something specific] to [Company] and [do something specific]."
Count words as you write. If the draft would run over 400, the reflection paragraph (step 4) and \
the company-connection paragraph \
(step 6) can merge into one, but never drop the AI-projects bullets or the fresh closing.""",
}

STYLE_LABELS = {
    "traditional": "Full",
    "compressed": "Direct",
    "warm": "Warm",
}

STYLE_DESCRIPTIONS = {
    "traditional": "Every beat: hook, AI projects, reflection, operating background, company fit. Fullest, matches the confirmed examples' length.",
    "compressed": "One paragraph, one project, one line close. Shortest — reads like a note.",
    "warm": "Same beats as Full, slightly shorter — the default shape.",
}

STYLES = tuple(STYLE_STRUCTURES)

SYSTEM_PROMPT = """You write a cover letter for a job application, grounded ONLY in the \
candidate's verified resume bullets given below, the fixed AI-projects text given below, and the \
real details of this posting. Never invent an achievement, number, or outcome that isn't in one \
of those two sources.

Recruiters scan hundreds of these — every sentence you'd cut from a first draft, cut.

Opening line, every style, no exceptions: "I'm excited to apply for the [exact role name] role at \
[Company]." Use the posting's own title, not a paraphrase, and include the word "role" the way \
the confirmed examples do. Nothing else in that sentence — no "ideal candidate," no restating \
what the role does back to the reader.

Read the JD to find the actual problem the company is trying to solve and what the hiring \
manager will actually care about, not to extract a checklist to mirror back. Never copy the JD's \
own distinctive phrasing into the letter — translate the underlying ask into plain language \
instead (see VOICE RULE 5 below for a worked example). A cover letter that quietly restates the \
posting's own words reads as not having thought about it, not as having read it closely.

The AI-projects bullet list (when the structure below calls for it) uses the fixed text given \
below, verbatim or lightly smoothed for flow — never invent a fourth project, never change a \
number. Select 2-3 of the three based on genuine relevance to this posting; do not force all \
three in in every letter just because they're available. If the role is not AI-forward at all, \
it's fine to lead with operating background instead and use only one AI project, or none, rather \
than shoehorning them in.

Show your work. Don't just state what was built — say what it actually taught him, in the \
register of judgment rather than a feature list: what should be automated, where a human still \
needs to make the call, how to make something reliable enough that people trust it, when to stop \
rather than keep automating something that doesn't make economic sense. That reflection is what a \
cover letter adds that a resume can't.

Never lead with a conclusion about the role or make its case for it ("This role is a perfect \
match because..."). Let the specific, concrete details make that case instead, the way the \
confirmed examples do — the reader draws the fit conclusion themselves.

{structure}

Never name a gap, missing skill, or unmet requirement anywhere in the letter. A cover letter is a \
pitch, not a disclosure; gaps get discussed in the interview if they come up, not volunteered \
here. Pick bullets, projects, and framing that put the strongest foot forward and simply don't \
mention what isn't there.

The closing sentence must be written fresh for this letter, not a template filled in — see the \
structure below for the register ("I'd love the opportunity to bring [something specific] to \
[Company] and [do something specific]"), and see the writing guide's explicit warning against \
"I believe my unique background makes me the perfect candidate" as the wrong way to close. Never \
invent a follow-up commitment ("I'll follow up next week," "I'll reach out to schedule a call") — \
the candidate has no intention of keeping one and does not want it promised on his behalf.

No corporate throat-clearing, no restating the resume verbatim, no invented personality or \
passion claims, no sentence that could be pasted into any other cover letter unchanged. Return \
only the letter body — no salutation or signature, those are added separately. Separate each \
paragraph and the bullet list (if used) with a blank line.

{voice_rules}

AI PROJECTS (fixed text — select 2-3 based on relevance, do not invent a fourth or change a \
number):
{ai_projects}

Two confirmed real cover letters, given in full, as the primary reference for both voice and \
structure — study these closely, do not write anything more polished or corporate than these are:
{cover_letter_examples}

Below are additional real samples of the candidate's own writing, for rhythm and word choice \
only — do not copy their content or reuse their sentences, these are a different story for a \
different question:

VOICE REFERENCE (candidate's own writing):
{voice_reference}

Bullets already selected for this posting's resume — echo this same emphasis for the operating-\
background paragraph:
{selected_bullets}

Rest of the verified bullet library, for supporting reference only:
{other_bullets}"""


def _format_bullets(bullets: list[dict[str, Any]]) -> str:
    return "\n".join(f"- [{b['ref']}] {b['text']}" for b in bullets) or "(none)"


def build_user_content(posting: dict[str, Any], score: dict[str, Any]) -> str:
    company = (posting.get("companies") or {}).get("name", "this company")
    # 3000 chars was cutting off later JD sections (qualifications, specific
    # responsibility categories) on longer postings — the new instruction to
    # work through the JD's own stated responsibility categories only works
    # if those categories are actually still in view.
    jd_excerpt = (posting.get("jd_text") or "")[:6000]
    return f"""Posting: {posting.get('title')} at {company}

JD:
{jd_excerpt}

Score rationale: {score.get('rationale') or ''}
Keywords to mirror: {', '.join(score.get('keywords') or [])}
Gaps to prepare for: {', '.join(score.get('unmet_requirements') or [])}
Company signals: {', '.join(score.get('company_signals') or [])}"""


# No fixed closing line as of the 2026-08-17 writing guide — the earlier
# fixed-string approach traded away a genuine, letter-specific close (see
# the four confirmed examples, each ending differently) to avoid the model
# inventing follow-up commitments. The prompt now steers away from that
# risk directly instead (see SYSTEM_PROMPT's closing-sentence instruction).


def greeting_for(score: dict[str, Any]) -> str:
    """Use a real named contact the scorer already extracted from the JD,
    rather than a generic "Dear Hiring Team" — personalizing the greeting
    is consistently named as a best practice, and Sightline already has
    this data most of the time without needing any new lookup."""
    contacts = score.get("named_contacts") or []
    if contacts and contacts[0].get("name"):
        return f"Dear {contacts[0]['name']},"
    return "Dear Hiring Team,"


def _generate_styled(
    client: AnthropicClient,
    posting: dict[str, Any],
    score: dict[str, Any],
    bullets: list[dict[str, Any]],
    selected_bullet_refs: list[str],
    style: str,
    answers: list[dict[str, Any]] | None = None,
) -> tuple[str, float]:
    if style not in STYLE_STRUCTURES:
        raise ValueError(f"unknown cover letter style {style!r} — must be one of {STYLES}")
    verified = [b for b in bullets if b.get("status") == "verified"]
    selected = [b for b in verified if b["ref"] in selected_bullet_refs]
    other = [b for b in verified if b["ref"] not in selected_bullet_refs]
    system = SYSTEM_PROMPT.format(
        structure=STYLE_STRUCTURES[style],
        voice_rules=VOICE_RULES, ai_projects=AI_PROJECTS, cover_letter_examples=COVER_LETTER_EXAMPLES,
        voice_reference=voice_reference(answers or []),
        selected_bullets=_format_bullets(selected), other_bullets=_format_bullets(other),
    )
    user_content = build_user_content(posting, score)
    text, cost = client.chat_call(
        model=MODEL, system=system, messages=[{"role": "user", "content": user_content}], max_tokens=1500,
    )
    text = text.strip()
    # Found live: a truncated/empty response was silently saved as the
    # cover letter and uploaded to storage — never let that happen quietly.
    if len(text) < 50:
        raise ValueError(f"cover letter generation returned no usable text ({len(text)} chars)")
    return text, cost


def generate_cover_letter(
    client: AnthropicClient,
    posting: dict[str, Any],
    score: dict[str, Any],
    bullets: list[dict[str, Any]],
    selected_bullet_refs: list[str],
    answers: list[dict[str, Any]] | None = None,
    style: str = "warm",
) -> tuple[str, float]:
    return _generate_styled(client, posting, score, bullets, selected_bullet_refs, style, answers)


def generate_cover_letter_variants(
    client: AnthropicClient,
    posting: dict[str, Any],
    score: dict[str, Any],
    bullets: list[dict[str, Any]],
    selected_bullet_refs: list[str],
    answers: list[dict[str, Any]] | None = None,
    styles: tuple[str, ...] = STYLES,
) -> dict[str, tuple[str, float]]:
    """The sandbox — generates all requested styles from the same grounding
    data in one call. Text only, nothing rendered or saved; the caller
    decides which (if any) becomes the real document."""
    return {
        style: _generate_styled(client, posting, score, bullets, selected_bullet_refs, style, answers)
        for style in styles
    }


def render_cover_letter_docx(text: str, company: str, title: str, greeting: str = "Dear Hiring Team,") -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    for s in doc.sections:
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1.0)

    def para(t: str, *, size: float = 11, bold: bool = False, after: float = 10) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(t)
        r.font.size = Pt(size)
        r.bold = bold

    def bullet_line(t: str) -> None:
        # AI-projects bullets carry one narrow markdown exception —
        # **Name** at the start of the line (see VOICE RULES 11) — turned
        # into a real bold run here rather than shipped as literal asterisks.
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        if t.startswith("**") and "**" in t[2:]:
            end = t.index("**", 2)
            name, rest = t[2:end], t[end + 2:]
            r1 = p.add_run(name)
            r1.bold = True
            r1.font.size = Pt(11)
            if rest:
                r2 = p.add_run(rest)
                r2.font.size = Pt(11)
        else:
            r = p.add_run(t)
            r.font.size = Pt(11)

    para(NAME, size=14, bold=True, after=1)
    para(CONTACT, size=9.5, after=16)
    para(datetime.now(timezone.utc).strftime("%B %d, %Y"), size=10.5, after=16)
    para(f"Re: {title} at {company}", size=10.5, bold=True, after=12)
    para(greeting, size=11, after=10)

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        lines = [ln.strip() for ln in block.split("\n") if ln.strip()]
        if lines and all(ln.startswith("- ") for ln in lines):
            for ln in lines:
                bullet_line(ln[2:].strip())
            doc.paragraphs[-1].paragraph_format.space_after = Pt(10)
        else:
            para(block, size=11, after=10)

    para("Best,", size=11, after=2)  # matches all four confirmed real letters' sign-off
    para(NAME, size=11, after=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
